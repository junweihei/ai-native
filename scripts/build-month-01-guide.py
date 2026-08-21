from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "第1月_完整执行指导手册_V2.0.md"
OUTPUT = ROOT / "AI_Native_第1月完整执行指导手册_V2.0.docx"

INK = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHTER_BLUE = "F4F7FA"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
GOLD = "A06A00"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_fonts(run, latin="Calibri", east_asia="Microsoft YaHei", size=None,
              bold=None, italic=None, color=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "left" if edge == "start" else "right" if edge == "end" else edge
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D7DEE8", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def compute_widths(rows):
    cols = len(rows[0])
    max_chars = []
    for idx in range(cols):
        length = max(len(re.sub(r"\[([^]]+)\]\([^)]+\)", r"\1", row[idx])) for row in rows)
        max_chars.append(max(6, min(length, 44)))
    if cols == 2:
        max_chars[0] = min(max_chars[0], 18)
    total = sum(max_chars)
    minimum = 1100 if cols >= 4 else 1400 if cols == 3 else 1800
    widths = [max(minimum, round(CONTENT_DXA * x / total)) for x in max_chars]
    diff = CONTENT_DXA - sum(widths)
    widths[-1] += diff
    if widths[-1] < minimum:
        short = minimum - widths[-1]
        donor = max(range(cols - 1), key=lambda i: widths[i])
        widths[donor] -= short
        widths[-1] = minimum
    return widths


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.width = Inches(widths[idx] / 1440)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    rpr.extend([rfonts, color, underline, size])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([rpr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\[[^\]]+\]\([^)]+\)|`[^`]+`|\*\*[^*]+\*\*)")


def add_inline(paragraph, text, size=11, bold=False, color=None):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_fonts(run, size=size, bold=bold, color=color)
        token = match.group(0)
        if token.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            add_hyperlink(paragraph, m.group(1), m.group(2))
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_fonts(run, latin="Consolas", east_asia="Microsoft YaHei", size=max(9.5, size - 0.5), color=DARK_BLUE)
            run.font.highlight_color = None
        else:
            run = paragraph.add_run(token[2:-2])
            set_fonts(run, size=size, bold=True, color=color)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_fonts(run, size=size, bold=bold, color=color)


def add_num_definition(doc, fmt, text, left, hanging, font="Calibri"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.extend([tabs, ind, spacing])
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rpr.append(rfonts)
    level.extend([start, num_fmt, lvl_text, lvl_jc, ppr, rpr])
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, numid])


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.text = "AI Native  ·  第 1 月完整执行指导"
    set_paragraph_spacing(p, after=0, line=1.0)
    for run in p.runs:
        set_fonts(run, size=8.5, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, after=0, line=1.0)
    run = p.add_run("V2.0  |  ")
    set_fonts(run, size=8.5, color=MUTED)
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])
    set_fonts(run, size=8.5, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=84, after=16, line=1.0)
    run = p.add_run("AI NATIVE 学习手册")
    set_fonts(run, size=11, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=10, line=1.0)
    run = p.add_run("第 1 月完整执行指导")
    set_fonts(run, size=30, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=28, line=1.2)
    run = p.add_run("图 — 拆 — 练 — 用 — 证 — 迭代")
    set_fonts(run, size=16, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=8, line=1.25)
    run = p.add_run("4 周 · 20 个学习日 · 每天 60–90 分钟")
    set_fonts(run, size=12, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=62, line=1.25)
    run = p.add_run("从整体认知到一个可验证的真实问题")
    set_fonts(run, size=11, italic=True, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=6, line=1.0)
    run = p.add_run("版本 V2.0")
    set_fonts(run, size=10, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0, line=1.0)
    run = p.add_run("资源核验：2026-08-21")
    set_fonts(run, size=9.5, color=MUTED)
    doc.add_page_break()


def add_contents(doc):
    p = doc.add_paragraph("使用说明", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p = doc.add_paragraph()
    add_inline(p, "这是一份执行手册。每天只打开当天卡片和指定成果文件；先闭卷，再看资料；达到验收标准即停止。")
    set_paragraph_spacing(p, after=10)

    call = doc.add_paragraph()
    set_paragraph_spacing(call, before=4, after=14, line=1.25)
    ppr = call._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHTER_BLUE)
    ppr.append(shd)
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    pbdr.append(left)
    ppr.append(pbdr)
    add_inline(call, "执行入口：工程根目录的“第1月_开始这里.md”。完整成果必须保存在 01-map 到 05-evidence，不能只留在聊天记录里。", bold=True, color=INK)

    p = doc.add_paragraph("目录", style="Heading 2")
    sections = [
        "一、现有路径审计结论", "二、本月目标与非目标", "三、月末六组成果",
        "四、每天固定执行法", "五、权威学习资源", "六、第 1 周：图",
        "七、第 2 周：拆", "八、第 3 周：练", "九、第 4 周：用、证、迭代",
        "十、快速反馈", "十一、中断处理", "十二、月末总验收", "十三、最简每日记录",
    ]
    bullet_id = add_num_definition(doc, "bullet", "•", 540, 270, "Symbol")
    for item in sections:
        p = doc.add_paragraph()
        apply_numbering(p, bullet_id)
        add_inline(p, item)
        set_paragraph_spacing(p, after=3)
    doc.add_page_break()


def add_callout(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=10, line=1.25)
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHTER_BLUE)
    ppr.append(shd)
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    pbdr.append(left)
    ppr.append(pbdr)
    add_inline(p, text, bold=True, color=INK)


def add_markdown_table(doc, rows):
    widths = compute_widths(rows)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (c_idx == 0 and len(row) > 2) else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, after=0, line=1.15)
            add_inline(p, value, size=9.5, bold=(r_idx == 0), color=(WHITE if r_idx == 0 else None))
            if r_idx == 0:
                shade_cell(cell, BLUE)
            elif r_idx % 2 == 0:
                shade_cell(cell, LIGHT_GRAY)
        if r_idx == 0:
            set_repeat_header(table.rows[0])
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=4, line=1.0)


def render_markdown(doc, text):
    lines = text.splitlines()
    # The cover supplies the source title and metadata.
    start = 0
    for idx, line in enumerate(lines):
        if line.startswith("## 一、"):
            start = idx
            break
    lines = lines[start:]
    bullet_id = add_num_definition(doc, "bullet", "•", 540, 270, "Symbol")
    decimal_id = add_num_definition(doc, "decimal", "%1.", 540, 270)
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        if not raw:
            i += 1
            continue
        if raw.strip() == "---":
            i += 1
            continue
        if raw.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[i + 1]):
            table_lines = [raw]
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            rows = []
            for line in table_lines:
                cells = [c.strip() for c in line.strip().strip("|").split("|")]
                rows.append(cells)
            add_markdown_table(doc, rows)
            continue
        if raw.startswith("### "):
            p = doc.add_paragraph(raw[4:].strip(), style="Heading 2")
            p.paragraph_format.keep_with_next = True
        elif raw.startswith("## "):
            title = raw[3:].strip()
            p = doc.add_paragraph(title, style="Heading 1")
        elif raw.startswith("# "):
            p = doc.add_paragraph(raw[2:].strip(), style="Heading 1")
        elif raw.startswith("> "):
            add_callout(doc, raw[2:].strip())
        elif re.match(r"^- ", raw):
            p = doc.add_paragraph()
            apply_numbering(p, bullet_id)
            add_inline(p, raw[2:].strip())
            set_paragraph_spacing(p, after=4)
        elif re.match(r"^\d+\. ", raw):
            p = doc.add_paragraph()
            apply_numbering(p, decimal_id)
            add_inline(p, re.sub(r"^\d+\.\s+", "", raw))
            set_paragraph_spacing(p, after=4)
        else:
            p = doc.add_paragraph()
            add_inline(p, raw)
            set_paragraph_spacing(p)
        i += 1


def build():
    doc = Document()
    style_document(doc)
    add_header_footer(doc.sections[0])
    doc.core_properties.title = "AI Native 第 1 月完整执行指导手册 V2.0"
    doc.core_properties.subject = "图—拆—练—用—证—迭代的 20 日执行路径"
    doc.core_properties.author = "AI Native Learning Project"
    doc.core_properties.keywords = "AI Native, 学习规划, Agent, 评测, 案例拆解"
    add_cover(doc)
    add_contents(doc)
    render_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
